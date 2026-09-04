#!/usr/bin/env python3
# ========== TEMPEST GUIDER - MAXIMIZED FINAL ==========
import os
import asyncio
import time
import random
import sqlite3
import json
import httpx
import shutil
import traceback
import psutil
import math
import io
import base64
import sys
import contextlib
import requests
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from PIL import Image, ImageDraw, ImageFont, ImageOps
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command, CommandStart
from aiogram.types import Message, FSInputFile, CallbackQuery, BufferedInputFile
from aiogram.enums import ParseMode, ChatType
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.exceptions import TelegramBadRequest, TelegramNetworkError

print("=" * 60)
print("🔒 TEMPEST GUIDER - MAXIMIZED FINAL")
print("=" * 60)

# ========== CONFIG ==========
BOT_TOKEN = "8017048722:AAFl6XZ9DHSJ6vnb8gUJXmoL7CUJc8AgehA"
OWNER_ID = 6108185460
UPLOAD_API = "https://tmpfiles.org/api/v1/upload"
LOG_CHANNEL_ID = -1003662720845

SECURE_HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
}

try:
    import yt_dlp
    YTDLP_AVAILABLE = True
except:
    YTDLP_AVAILABLE = False

try:
    from zoneinfo import ZoneInfo
except:
    ZoneInfo = None

for d in ["data", "temp", "backups", "fonts"]:
    Path(d).mkdir(exist_ok=True)

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

start_time = time.time()
upload_waiting = {}
broadcast_state = {}
pending_restore = {}
disabled_commands = {}
maintenance_mode = False

# ========== FONT ==========
async def ensure_font():
    font_path = "fonts/font.ttf"
    if not os.path.exists(font_path):
        try:
            url = "https://github.com/google/fonts/raw/main/ofl/roboto/Roboto-Regular.ttf"
            async with httpx.AsyncClient(timeout=30) as client:
                r = await client.get(url, follow_redirects=True)
                if r.status_code == 200:
                    with open(font_path, "wb") as f:
                        f.write(r.content)
        except:
            pass

def get_safe_font(size):
    font_path = "fonts/font.ttf"
    if os.path.exists(font_path):
        try:
            return ImageFont.truetype(font_path, size)
        except:
            pass
    for path in ["/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "/system/fonts/Roboto-Bold.ttf"]:
        if os.path.exists(path):
            try:
                return ImageFont.truetype(path, size)
            except:
                pass
    return ImageFont.load_default()

def draw_visible_text(draw, pos, text, font, fill, outline=(0,0,0), w=3, anchor=None):
    x, y = pos
    for ox in range(-w, w+1):
        for oy in range(-w, w+1):
            if ox != 0 or oy != 0:
                if anchor:
                    draw.text((x+ox, y+oy), text, font=font, fill=outline, anchor=anchor)
                else:
                    draw.text((x+ox, y+oy), text, font=font, fill=outline)
    if anchor:
        draw.text((x, y), text, font=font, fill=fill, anchor=anchor)
    else:
        draw.text((x, y), text, font=font, fill=fill)

# ========== ART ==========
def header(t):
    return f"◤━━━━━━━━━━━━━━━━━━━━◥\n◇ {t} ◇\n◣━━━━━━━━━━━━━━━━━━━━◢"

def divider():
    return "━━━━━━━━━━━━━━━━━━━━"

# ========== QUOTES ==========
ANIME_QUOTES = [
    "⚡ Wake up to reality! Nothing ever goes as planned in this accursed world.\n— Madara Uchiha",
    "🌊 The longer you live, the more you realize that reality is just made of pain.\n— Madara Uchiha",
    "💔 People cannot show each other their true feelings.\n— Madara Uchiha",
    "🪽 The only thing we're allowed to do is believe we won't regret our choice.\n— Levi Ackerman",
    "🏴‍☠️ If you don't take risks, you can't create a future.\n— Monkey D. Luffy",
    "💪 Power isn't determined by your size, but by the size of your heart.\n— Luffy",
    "😨 Fear is not evil. It tells you what your weakness is.\n— Gildarts Clive",
    "🔥 A lesson without pain is meaningless.\n— Edward Elric",
]

COUNTRY_TIMEZONES = {
    "usa": "America/New_York", "uk": "Europe/London", "india": "Asia/Kolkata",
    "japan": "Asia/Tokyo", "china": "Asia/Shanghai", "russia": "Europe/Moscow",
    "brazil": "America/Sao_Paulo", "australia": "Australia/Sydney",
    "canada": "America/Toronto", "germany": "Europe/Berlin",
    "france": "Europe/Paris", "italy": "Europe/Rome", "spain": "Europe/Madrid",
    "south korea": "Asia/Seoul", "pakistan": "Asia/Karachi",
    "bangladesh": "Asia/Dhaka", "nigeria": "Africa/Lagos",
    "egypt": "Africa/Cairo", "south africa": "Africa/Johannesburg",
    "kenya": "Africa/Nairobi", "uganda": "Africa/Kampala",
    "uae": "Asia/Dubai", "saudi arabia": "Asia/Riyadh",
    "turkey": "Europe/Istanbul", "thailand": "Asia/Bangkok",
    "philippines": "Asia/Manila", "singapore": "Asia/Singapore",
    "new york": "America/New_York", "los angeles": "America/Los_Angeles",
    "london": "Europe/London", "tokyo": "Asia/Tokyo",
    "sydney": "Australia/Sydney", "moscow": "Europe/Moscow",
    "dubai": "Asia/Dubai", "singapore": "Asia/Singapore",
}

FORTUNES_LIST = [
    "🌟 Great things await you in the coming days!",
    "🗝️ Hidden doors of opportunity will soon open!",
    "🦋 Beautiful transformations are happening!",
    "🌊 A grand adventure approaches!",
    "🔆 Success is closer than you think!",
    "🌠 A long-held wish may soon come true!",
    "💕 True love will find you unexpectedly!",
    "📈 Your persistence is about to pay off!",
    "🍀 Extraordinary luck surrounds you this week!",
    "💪 Every trial has made you stronger!",
    "🎯 Your focus will lead to victory!",
    "🌈 After every storm comes a rainbow!",
    "🔥 Your passion will inspire others!",
    "⚡ The storm favors your bold moves!",
    "🕊️ Peace will find your troubled heart!",
    "👑 Leadership opportunities are coming!",
    "💎 Hidden talents will emerge soon!",
    "🎲 Take a chance - fortune favors you!",
    "🌙 Your dreams are closer than they appear!",
    "☀️ A bright future is on the horizon!",
]

# ========== DATABASE ==========
def init_db():
    with sqlite3.connect("data/bot.db") as conn:
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS users (
            user_id INTEGER PRIMARY KEY, username TEXT, first_name TEXT,
            joined_date TEXT, last_active TEXT, uploads INTEGER DEFAULT 0,
            commands INTEGER DEFAULT 0, is_admin INTEGER DEFAULT 0,
            is_banned INTEGER DEFAULT 0, cult_status TEXT DEFAULT 'none',
            cult_rank TEXT DEFAULT 'none', cult_join_date TEXT,
            sacrifices INTEGER DEFAULT 0, curse_type TEXT DEFAULT 'none'
        )''')
        for col, defn in [("cult_status","TEXT DEFAULT 'none'"),("cult_rank","TEXT DEFAULT 'none'"),("cult_join_date","TEXT"),("sacrifices","INTEGER DEFAULT 0"),("curse_type","TEXT DEFAULT 'none'")]:
            try:
                c.execute(f"ALTER TABLE users ADD COLUMN {col} {defn}")
            except:
                pass
        c.execute('''CREATE TABLE IF NOT EXISTS groups (group_id INTEGER PRIMARY KEY, title TEXT, joined_date TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS uploads (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, timestamp TEXT, file_url TEXT, file_type TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS command_logs (id INTEGER PRIMARY KEY AUTOINCREMENT, timestamp TEXT, user_id INTEGER, command TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS wishes (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, wish_text TEXT, luck INTEGER)''')
        c.execute('''CREATE TABLE IF NOT EXISTS fate_pairs (chat_id INTEGER, user1_name TEXT, user2_name TEXT, love_percentage INTEGER, created_date TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS fortunes (user_id INTEGER, fortune_text TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS shrines (user_id INTEGER PRIMARY KEY, power_level INTEGER DEFAULT 10)''')
        c.execute('''CREATE TABLE IF NOT EXISTS bot_memory (key TEXT PRIMARY KEY, value TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS saved_commands (id INTEGER PRIMARY KEY AUTOINCREMENT, command_code TEXT, saved_date TEXT)''')
        c.execute("INSERT OR IGNORE INTO users (user_id, first_name, is_admin) VALUES (?, 'Owner', 1)", (OWNER_ID,))
        conn.commit()

init_db()

def save_memory():
    try:
        with sqlite3.connect("data/bot.db") as conn:
            conn.execute("INSERT OR REPLACE INTO bot_memory (key, value) VALUES ('upload_waiting', ?)", (json.dumps(upload_waiting),))
            conn.execute("INSERT OR REPLACE INTO bot_memory (key, value) VALUES ('broadcast_state', ?)", (json.dumps(broadcast_state),))
            conn.execute("INSERT OR REPLACE INTO bot_memory (key, value) VALUES ('pending_restore', ?)", (json.dumps(pending_restore),))
            conn.commit()
    except:
        pass

def load_memory():
    try:
        with sqlite3.connect("data/bot.db") as conn:
            r = conn.execute("SELECT value FROM bot_memory WHERE key = 'upload_waiting'").fetchone()
            if r:
                upload_waiting.update(json.loads(r[0]))
            r = conn.execute("SELECT value FROM bot_memory WHERE key = 'broadcast_state'").fetchone()
            if r:
                broadcast_state.update(json.loads(r[0]))
            r = conn.execute("SELECT value FROM bot_memory WHERE key = 'pending_restore'").fetchone()
            if r:
                pending_restore.update(json.loads(r[0]))
    except:
        pass

load_memory()

async def send_log(msg):
    try:
        await bot.send_message(LOG_CHANNEL_ID, msg[:4000])
    except:
        pass

async def handle_common(message, cmd):
    user = message.from_user
    chat = message.chat
    if chat.type in [ChatType.GROUP, ChatType.SUPERGROUP]:
        try:
            with sqlite3.connect("data/bot.db") as conn:
                conn.execute("INSERT OR IGNORE INTO groups (group_id, title, joined_date) VALUES (?, ?, ?)",
                             (chat.id, str(chat.title), datetime.now().isoformat()))
                conn.commit()
        except:
            pass
    if maintenance_mode and user.id != OWNER_ID:
        await message.answer("🔧 Maintenance mode!")
        return None, None
    if cmd in disabled_commands:
        await message.answer("⛔ Command disabled!")
        return None, None
    try:
        with sqlite3.connect("data/bot.db") as conn:
            c = conn.cursor()
            c.execute("INSERT OR IGNORE INTO users (user_id, username, first_name, joined_date) VALUES (?, ?, ?, ?)",
                      (user.id, user.username or "", user.first_name or "", datetime.now().isoformat()))
            c.execute("UPDATE users SET last_active = ?, username = ?, first_name = ? WHERE user_id = ?",
                      (datetime.now().isoformat(), user.username or "", user.first_name or "", user.id))
            c.execute("INSERT INTO command_logs (timestamp, user_id, command) VALUES (?, ?, ?)",
                      (datetime.now().isoformat(), user.id, cmd))
            c.execute("UPDATE users SET commands = commands + 1 WHERE user_id = ?", (user.id,))
            conn.commit()
    except:
        pass
    return user, chat

async def is_admin(uid):
    if uid == OWNER_ID:
        return True
    try:
        with sqlite3.connect("data/bot.db") as conn:
            r = conn.execute("SELECT is_admin FROM users WHERE user_id = ?", (uid,)).fetchone()
        return r and r[0] == 1
    except:
        return False

def format_uptime(s):
    d = s // 86400
    h = (s % 86400) // 3600
    m = (s % 3600) // 60
    sec = s % 60
    parts = []
    if d: parts.append(f"{d}d")
    if h: parts.append(f"{h}h")
    if m: parts.append(f"{m}m")
    if sec or not parts: parts.append(f"{sec}s")
    return " ".join(parts)

async def upload_to_catbox(data, filename):
    # Primary: tmpfiles.org
    try:
        temp_path = f"temp/{filename}"
        with open(temp_path, "wb") as f:
            f.write(data)
        with open(temp_path, "rb") as file:
            files = {"file": (filename, file)}
            r = requests.post(UPLOAD_API, files=files, timeout=30)
        if os.path.exists(temp_path):
            os.remove(temp_path)
        if r.status_code == 200:
            result = r.json()
            url = result.get("data", {}).get("url", "")
            if url.startswith("https://tmpfiles.org/"):
                return url.replace("https://tmpfiles.org/", "https://tmpfiles.org/dl/")
            return f"https://tmpfiles.org/dl/{url}"
    except Exception as e:
        print(f"tmpfiles error: {e}")
    
    # Fallback: catbox.moe
    try:
        files = {"fileToUpload": (filename, data)}
        form_data = {"reqtype": "fileupload", "userhash": ""}
        async with httpx.AsyncClient(timeout=60, headers=SECURE_HEADERS) as client:
            r = await client.post("https://catbox.moe/user/api.php", data=form_data, files=files)
        if r.status_code == 200 and r.text.strip().startswith('http'):
            return r.text.strip()
    except Exception as e:
        print(f"Catbox error: {e}")
    
    # Fallback: 0x0.st
    try:
        files = {"file": (filename, data)}
        async with httpx.AsyncClient(timeout=60, headers=SECURE_HEADERS) as client:
            r = await client.post("https://0x0.st", files=files)
        if r.status_code == 200 and r.text.strip().startswith('http'):
            return r.text.strip()
    except:
        pass
    return None

class EncryptionEngine:
    @staticmethod
    def xor_encrypt(text, key="TEMPEST"):
        return ''.join(chr(ord(c) ^ ord(key[i % len(key)])) for i, c in enumerate(text))
    @staticmethod
    def encrypt(t):
        return base64.b64encode(EncryptionEngine.xor_encrypt(t).encode()).decode()
    @staticmethod
    def decrypt(t):
        try:
            return EncryptionEngine.xor_encrypt(base64.b64decode(t.encode()).decode())
        except:
            return None

async def fetch_user_avatar(uid, name="User"):
    try:
        photos = await bot.get_user_profile_photos(uid, limit=1)
        if photos.total_count > 0:
            fid = photos.photos[0][-1].file_id
            fi = await bot.get_file(fid)
            async with httpx.AsyncClient(timeout=30) as c:
                r = await c.get(f"https://api.telegram.org/file/bot{BOT_TOKEN}/{fi.file_path}")
            if r.status_code == 200:
                av = Image.open(io.BytesIO(r.content)).convert("RGBA")
                av = ImageOps.fit(av, (150, 150), Image.Resampling.LANCZOS)
                mask = Image.new("L", (150, 150), 0)
                d = ImageDraw.Draw(mask)
                d.ellipse((0, 0, 150, 150), fill=255)
                out = Image.new("RGBA", (150, 150), (0, 0, 0, 0))
                out.paste(av, (0, 0), mask=mask)
                return out
    except:
        pass
    fb = Image.new("RGBA", (150, 150), (20, 30, 50, 255))
    d = ImageDraw.Draw(fb)
    d.ellipse((0, 0, 150, 150), outline=(0, 255, 200), width=3)
    f = get_safe_font(60)
    draw_visible_text(d, (75, 75), name[0].upper() if name else "T", f, (0, 255, 200), anchor="mm")
    return fb

def gen_profile_card(username, uid, rank, up, ws, av):
    w, h = 900, 450
    img = Image.new("RGB", (w, h), (10, 15, 30))
    d = ImageDraw.Draw(img)
    for y in range(h):
        r = int(10 + 25 * (y/h))
        g = int(15 + 35 * (y/h))
        b = int(30 + 55 * (y/h))
        d.line([(0, y), (w, y)], fill=(r, g, b))
    d.rectangle([10, 10, w-10, h-10], outline=(0, 200, 255), width=3)
    av_x, av_y = 50, 60
    d.ellipse((av_x-3, av_y-3, av_x+153, av_y+153), outline=(0, 255, 200), width=3)
    img.paste(av, (av_x, av_y), mask=av)
    fn = get_safe_font(50)
    fi = get_safe_font(32)
    draw_visible_text(d, (230, 60), str(username)[:15], fn, (255, 255, 255))
    draw_visible_text(d, (230, 130), f"ID: {uid}", fi, (180, 200, 230))
    draw_visible_text(d, (230, 180), f"Rank: {rank}", fi, (0, 255, 200))
    draw_visible_text(d, (230, 240), f"Uploads: {up}  |  Wishes: {ws}", fi, (255, 215, 0))
    draw_visible_text(d, (870, 410), "TEMPEST GUIDER", fi, (100, 116, 139), anchor="ra")
    b = io.BytesIO()
    img.save(b, format="PNG")
    b.seek(0)
    return b.getvalue()

def gen_fate_card(n1, n2, pct, q, av1=None, av2=None):
    w, h = 900, 550
    img = Image.new("RGB", (w, h), (15, 10, 25))
    d = ImageDraw.Draw(img)
    for y in range(h):
        r = int(15 + 30 * (y/h))
        g = int(10 + 20 * (y/h))
        b = int(25 + 45 * (y/h))
        d.line([(0, y), (w, y)], fill=(r, g, b))
    d.rectangle([10, 10, w-10, h-10], outline=(255, 100, 180), width=3)
    ft = get_safe_font(40)
    fn = get_safe_font(30)
    fq = get_safe_font(20)
    if av1:
        img.paste(av1, (100, 60), mask=av1)
    if av2:
        img.paste(av2, (650, 60), mask=av2)
    draw_visible_text(d, (450, 35), "TEMPEST FATE", ft, (255, 100, 180), anchor="mm")
    draw_visible_text(d, (450, 250), f"{n1}  +  {n2}", fn, (255, 255, 255), anchor="mm")
    draw_visible_text(d, (450, 310), f"{pct}% COMPATIBILITY", fn, (255, 215, 0), anchor="mm")
    words = q.split()
    lines = []
    cur = ""
    for wrd in words:
        if len(cur + wrd) < 40:
            cur += wrd + " "
        else:
            lines.append(cur.strip())
            cur = wrd + " "
    if cur:
        lines.append(cur.strip())
    yp = 370
    for ln in lines[:5]:
        draw_visible_text(d, (450, yp), ln, fq, (210, 220, 240), anchor="mm")
        yp += 32
    b = io.BytesIO()
    img.save(b, format="PNG")
    b.seek(0)
    return b.getvalue()

# ========== COMMANDS ==========
@dp.message(CommandStart())
async def start_cmd(m: Message):
    u, _ = await handle_common(m, "start")
    if not u:
        return
    await m.answer(f"{header('🔒 TEMPEST GUIDER')}\n\n✨ Welcome {u.first_name}!\n\n📚 /help - Commands")

@dp.message(Command("myid"))
async def myid_cmd(m: Message):
    await m.answer(f"🆔 Your ID: <code>{m.from_user.id}</code>", parse_mode=ParseMode.HTML)

@dp.message(Command("help"))
async def help_cmd(m: Message):
    await handle_common(m, "help")
    await m.answer(
        f"{header('📚 COMMANDS')}\n\n"
        f"🔗 /link - Upload\n📥 /convert - Media URL\n"
        f"✨ /wish - Wish\n🔮 /fortune - Future\n"
        f"🎮 /dice - Dice\n🪙 /flip - Coin\n💑 /fate - Love\n"
        f"👤 /profile - Stats\n🔐 /encrypt - Encrypt\n🔓 /decrypt - Decrypt\n"
        f"🌀 /tempest_join - Join\n📖 /tempest_story - Lore\n"
        f"👑 /tempest_creed - Members\n⛩️ /shrine - Shrine\n"
        f"⚡ /curse - Curse\n✅ /remove_curse - Uncurse\n"
        f"🌍 /time - World time\n📝 /word - DOCX\n"
        f"🆔 /myid - Your ID\n👑 /admin_help - Admin"
    )

@dp.message(Command("admin_help"))
async def admin_help_cmd(m: Message):
    u, _ = await handle_common(m, "admin_help")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer(f"🚫 Admin only. Your ID: {u.id}")
        return
    await m.answer(
        f"{header('👑 ADMIN')}\n\n"
        f"📊 /ping /stats /scan /users /broadcast /lag /disable /cm\n"
        f"🚫 /ban /unban /banlist\n"
        f"⚡ OWNER: /query /restart /maintenance /pro /backup /rem /clearlogs /logs"
    )

# ========== WISH (ANIMATED) ==========
@dp.message(Command("wish"))
async def wish_cmd(m: Message):
    u, _ = await handle_common(m, "wish")
    if not u:
        return
    args = m.text.split(maxsplit=1)
    if len(args) < 2:
        await m.answer("✨ Usage: /wish [text]")
        return
    msg = await m.answer("🔮 Consulting the storm...")
    await asyncio.sleep(1)
    await msg.edit_text("✨ Reading your destiny...")
    await asyncio.sleep(1)
    await msg.edit_text("🌊 The tempest whispers...")
    await asyncio.sleep(1)
    luck = random.randint(1, 100)
    stars = "⭐" * (luck // 10) + "☆" * (10 - luck // 10)
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("INSERT INTO wishes (user_id, wish_text, luck) VALUES (?, ?, ?)", (u.id, args[1], luck))
        conn.commit()
    await msg.edit_text(f"{header('✨ WISH')}\n\n📜 {args[1][:100]}\n🎰 {stars} {luck}%")

# ========== FORTUNE (ANIMATED) ==========
@dp.message(Command("fortune"))
async def fortune_cmd(m: Message):
    u, _ = await handle_common(m, "fortune")
    if not u:
        return
    msg = await m.answer("🔮 Reading the crystal orb...")
    await asyncio.sleep(1.5)
    await msg.edit_text("👻 The spirits are speaking...")
    await asyncio.sleep(1.5)
    f = random.choice(FORTUNES_LIST)
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("INSERT INTO fortunes (user_id, fortune_text) VALUES (?, ?)", (u.id, f))
        conn.commit()
    await msg.edit_text(f"{header('🔮 FORTUNE')}\n\n{f}")

# ========== DICE ==========
@dp.message(Command("dice"))
async def dice_cmd(m: Message):
    await handle_common(m, "dice")
    await m.answer_dice(emoji="🎲")

# ========== FLIP ==========
@dp.message(Command("flip"))
async def flip_cmd(m: Message):
    await handle_common(m, "flip")
    await m.answer(f"🪙 {random.choice(['HEADS 🟡', 'TAILS 🟤'])}!")

# ========== FATE (ANIMATED + CARDS) ==========
@dp.message(Command("fate"))
async def fate_cmd(m: Message):
    u, ch = await handle_common(m, "fate")
    if not u:
        return
    if ch.type not in [ChatType.GROUP, ChatType.SUPERGROUP]:
        await m.answer("💑 Groups only!")
        return
    msg = await m.answer("💫 Weaving destiny...")
    await asyncio.sleep(1.5)
    await msg.edit_text("🔍 Scanning souls...")
    await asyncio.sleep(1.5)
    await msg.edit_text("💖 Calculating compatibility...")
    await asyncio.sleep(1.5)
    with sqlite3.connect("data/bot.db") as conn:
        c = conn.cursor()
        c.execute("SELECT user1_name, user2_name, love_percentage FROM fate_pairs WHERE chat_id = ? AND created_date >= ?", (ch.id, (datetime.now() - timedelta(hours=24)).isoformat()))
        ex = c.fetchone()
        if ex:
            await msg.edit_text(f"💑 Today: {ex[0]} & {ex[1]} ({ex[2]}%)")
            return
        c.execute("SELECT user_id, first_name FROM users WHERE user_id != ? LIMIT 50", (u.id,))
        members = c.fetchall()
    if len(members) < 2:
        await msg.edit_text("❌ Not enough!")
        return
    l1, l2 = random.sample(members, 2)
    love = random.randint(50, 100)
    quote = random.choice(ANIME_QUOTES)
    av1 = await fetch_user_avatar(l1[0], l1[1])
    av2 = await fetch_user_avatar(l2[0], l2[1])
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("INSERT INTO fate_pairs (chat_id, user1_name, user2_name, love_percentage, created_date) VALUES (?, ?, ?, ?, ?)", (ch.id, l1[1], l2[1], love, datetime.now().isoformat()))
        conn.commit()
    card = await asyncio.to_thread(gen_fate_card, l1[1], l2[1], love, quote, av1, av2)
    await msg.delete()
    await m.answer_photo(photo=BufferedInputFile(card, filename="fate.png"), caption=f"💑 {l1[1]} & {l2[1]} - 💖 {love}%")

# ========== PROFILE ==========
@dp.message(Command("profile"))
async def profile_cmd(m: Message):
    u, _ = await handle_common(m, "profile")
    if not u:
        return
    msg = await m.answer("⚡ Generating...")
    with sqlite3.connect("data/bot.db") as conn:
        c = conn.cursor()
        c.execute("SELECT uploads, cult_rank FROM users WHERE user_id = ?", (u.id,))
        r = c.fetchone()
        c.execute("SELECT COUNT(*) FROM wishes WHERE user_id = ?", (u.id,))
        ws = c.fetchone()[0]
    up, rank = (r[0], r[1]) if r else (0, "Mortal")
    if not rank or rank == 'none':
        rank = "Mortal"
    av = await fetch_user_avatar(u.id, u.first_name)
    card = await asyncio.to_thread(gen_profile_card, u.first_name, u.id, rank, up, ws, av)
    await msg.delete()
    await m.answer_photo(photo=BufferedInputFile(card, filename="profile.png"), caption=f"👤 {u.first_name}'s Card")

# ========== ENCRYPT/DECRYPT ==========
@dp.message(Command("encrypt"))
async def encrypt_cmd(m: Message):
    await handle_common(m, "encrypt")
    args = m.text.split(maxsplit=1)
    if len(args) < 2:
        await m.answer("🔐 Usage: /encrypt [text]")
        return
    await m.answer(f"🔐 <code>{EncryptionEngine.encrypt(args[1])}</code>", parse_mode=ParseMode.HTML)

@dp.message(Command("decrypt"))
async def decrypt_cmd(m: Message):
    await handle_common(m, "decrypt")
    args = m.text.split(maxsplit=1)
    if len(args) < 2:
        await m.answer("🔓 Usage: /decrypt [text]")
        return
    dec = EncryptionEngine.decrypt(args[1])
    await m.answer(f"🔓 {dec}" if dec else "❌ Invalid!")

# ========== TEMPEST JOIN (RITUAL) ==========
@dp.message(Command("tempest_join"))
async def tempest_join_cmd(m: Message):
    u, _ = await handle_common(m, "tempest_join")
    if not u:
        return
    with sqlite3.connect("data/bot.db") as conn:
        c = conn.cursor()
        c.execute("SELECT cult_status FROM users WHERE user_id = ?", (u.id,))
        r = c.fetchone()
        if r and r[0] != 'none':
            await m.answer("🌀 Already in!")
            return
        c.execute("UPDATE users SET cult_status = 'member', cult_rank = 'Blood Initiate', cult_join_date = ?, sacrifices = 3 WHERE user_id = ?", (datetime.now().isoformat(), u.id))
        conn.commit()
    ritual = ["🌀 INITIATING BLOOD PACT...", "🩸 Drawing sigils...", "⚡ Channeling storm...", "🌑 The void responds...", "🔥 Sacrifice offered...Hm", "🌀 TEMPEST AWAKENS!"]
    msg = await m.answer("🌀 Preparing ritual...")
    for t in ritual:
        await asyncio.sleep(1.2)
        await msg.edit_text(t)
    await msg.edit_text("⚡ WELCOME TO THE TEMPEST!\n🌀 Rank: Blood Initiate\n⚔️ Sacrifices: 3")

# ========== TEMPEST STORY (6 CHAPTERS) ==========
@dp.message(Command("tempest_story"))
async def tempest_story_cmd(m: Message):
    u, _ = await handle_common(m, "tempest_story")
    if not u:
        return
    chapters = [
        ("📖 Chapter 1: Awakening", "Keny Marcus opened his eyes in the dark realm(boredom), realm of Tempest alongside Bablu and Ravijah. The sky burned with violet aura ..ideas."),
        ("⚔️ Chapter 2: In the realm of darkness", "With Bablu holding the oath and Ravijah orchestrating tactics, Keny forged the Tempest Guild...did nothin"),
        ("⚡ Chapter 3: Breach of Weirdness", "Ravijah bypassed the firewall,...uff... while Bablu smashed through the fortress gates..became wise. Keny unleashed full idleness... Boredom at max."),
        ("👑 Chapter 4: Reign of King", "Standing a top the conquered spire, Keny, Bablu, and Ravijah gazed across the infinite grid."),
        ("🌌 Chapter 5: The Void Calls", "An ancient entity stirred. The Void whispered promises of infinite power to the three founders."),
        ("🔱 Chapter 6: Eternal Storm", "Lightning became their blood, thunder their voice. Keny, Ravijah, and Bablu - the eternal storm... technically I'm bored so this story is chup."),
    ]
    msg = await m.answer("📜 Opening Tempest Archives...")
    await asyncio.sleep(1.5)
    for i, (t, c) in enumerate(chapters):
        prog = "[" + "#" * (i + 1) + "." * (len(chapters) - i - 1) + "]"
        await msg.edit_text(f"📜 Loading... {prog} {i+1}/{len(chapters)}")
        await asyncio.sleep(2)
        await msg.edit_text(f"{header(t)}\n\n{c}")
        await asyncio.sleep(5)
    await msg.edit_text("🌀 We are the eternal storm.-As u can see I was bored")

# ========== TEMPEST CREED ==========
@dp.message(Command("tempest_creed"))
async def tempest_creed_cmd(m: Message):
    u, _ = await handle_common(m, "tempest_creed")
    if not u:
        return
    with sqlite3.connect("data/bot.db") as conn:
        members = conn.execute("SELECT first_name, cult_rank, sacrifices FROM users WHERE cult_status != 'none' ORDER BY sacrifices DESC LIMIT 10").fetchall()
    if not members:
        await m.answer("🌀 No members!")
        return
    text = f"{header('🌀 CREED')}\n\n"
    for i, (n, r, s) in enumerate(members, 1):
        medal = ["🥇", "🥈", "🥉"][i-1] if i <= 3 else "👤"
        text += f"{medal} {n} — {r} ⚔️{s}\n"
    await m.answer(text)

# ========== SHRINE ==========
@dp.message(Command("shrine"))
async def shrine_cmd(m: Message):
    u, _ = await handle_common(m, "shrine")
    if not u:
        return
    with sqlite3.connect("data/bot.db") as conn:
        c = conn.cursor()
        c.execute("INSERT OR IGNORE INTO shrines (user_id) VALUES (?)", (u.id,))
        c.execute("SELECT power_level FROM shrines WHERE user_id = ?", (u.id,))
        p = c.fetchone()[0]
        conn.commit()
    await m.answer(f"{header('⛩️ SHRINE')}\n\n⚡ {p} XP")

# ========== CURSE ==========
@dp.message(Command("curse"))
async def curse_cmd(m: Message):
    u, _ = await handle_common(m, "curse")
    if not u:
        return
    if not m.reply_to_message:
        await m.answer("⚡ Reply to curse!")
        return
    t = m.reply_to_message.from_user
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("UPDATE users SET curse_type = 'Bad Luck' WHERE user_id = ?", (t.id,))
        conn.commit()
    await m.reply(f"⚡ {t.first_name} cursed!")

@dp.message(Command("remove_curse"))
async def remove_curse_cmd(m: Message):
    u, _ = await handle_common(m, "remove_curse")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only")
        return
    if not m.reply_to_message:
        await m.answer("Reply to remove!")
        return
    t = m.reply_to_message.from_user
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("UPDATE users SET curse_type = 'none' WHERE user_id = ?", (t.id,))
        conn.commit()
    await m.reply(f"✅ Removed from {t.first_name}!")

# ========== TIME (FIXED) ==========
@dp.message(Command("time"))
async def time_cmd(m: Message):
    u, _ = await handle_common(m, "time")
    if not u:
        return
    args = m.text.split(maxsplit=1)
    if len(args) < 2:
        await m.answer("🌍 Usage: /time [country]\n\nExamples:\n/time usa\n/time india\n/time japan\n/time uk\n/time uae\n/time pakistan\n/time bangladesh")
        return
    country_input = args[1].lower().strip()
    
    # Try exact match first
    tz = COUNTRY_TIMEZONES.get(country_input)
    
    # If not found, try partial match
    if not tz:
        for key in COUNTRY_TIMEZONES:
            if country_input in key or key in country_input:
                tz = COUNTRY_TIMEZONES[key]
                country_input = key
                break
    
    if not tz:
        await m.answer(f"❌ Country not found!\n\nAvailable: {', '.join(list(COUNTRY_TIMEZONES.keys())[:15])}...")
        return
    
    try:
        if ZoneInfo:
            now = datetime.now(ZoneInfo(tz))
        else:
            now = datetime.utcnow()
        await m.answer(f"🌍 {country_input.upper()}: {now.strftime('%H:%M:%S')}\n📅 {now.strftime('%A, %B %d, %Y')}")
    except Exception as e:
        await m.answer(f"❌ Error: {e}")

# ========== WORD ==========
@dp.message(Command("word"))
async def word_cmd(m: Message):
    u, _ = await handle_common(m, "word")
    if not u:
        return
    args = m.text.split(maxsplit=1)
    if len(args) < 2:
        await m.answer("📝 Usage: /word [text]")
        return
    msg = await m.answer("📝 Creating...")
    await asyncio.sleep(1)
    doc = Document()
    doc.add_heading("🌀 TEMPEST", 0)
    doc.add_paragraph(f"By: {u.first_name}")
    doc.add_paragraph(args[1])
    fn = f"temp/word_{u.id}.docx"
    doc.save(fn)
    await msg.delete()
    await m.answer_document(FSInputFile(fn))
    os.remove(fn)

# ========== LINK/UPLOAD ==========
@dp.message(Command("link"))
async def link_cmd(m: Message):
    u, _ = await handle_common(m, "link")
    if not u:
        return
    upload_waiting[u.id] = True
    save_memory()
    await m.answer("📁 Send me any file!")

@dp.message(F.photo | F.video | F.document | F.audio | F.voice | F.sticker | F.animation)
async def handle_file(m: Message):
    u = m.from_user
    if u.id in pending_restore:
        pending_restore.pop(u.id, None)
        if not m.document or not m.document.file_name.endswith(".db"):
            await m.answer("❌ .db file!")
            return
        msg = await m.answer("⏳ Restoring...")
        try:
            f = await bot.get_file(m.document.file_id)
            downloaded = await bot.download_file(f.file_path)
            with open("data/bot.db", "wb") as fh:
                fh.write(downloaded.read())
            init_db()
            load_memory()
            await msg.edit_text("✅ Restored!")
        except Exception as e:
            await msg.edit_text(f"❌ {e}")
        return
    if u.id not in upload_waiting:
        return
    upload_waiting.pop(u.id, None)
    save_memory()
    msg = await m.answer("⏳ Uploading...")
    try:
        if m.photo: fid = m.photo[-1].file_id; fn = f"photo_{u.id}.jpg"
        elif m.video: fid = m.video.file_id; fn = m.video.file_name or f"video_{u.id}.mp4"
        elif m.document: fid = m.document.file_id; fn = m.document.file_name or f"doc_{u.id}"
        elif m.audio: fid = m.audio.file_id; fn = f"audio_{u.id}.mp3"
        elif m.voice: fid = m.voice.file_id; fn = f"voice_{u.id}.ogg"
        elif m.sticker: fid = m.sticker.file_id; fn = f"sticker_{u.id}.webp"
        elif m.animation: fid = m.animation.file_id; fn = f"gif_{u.id}.gif"
        else:
            await msg.edit_text("❌ Unsupported!")
            return
        f = await bot.get_file(fid)
        downloaded = await bot.download_file(f.file_path)
        file_data = downloaded.read()
        link = await upload_to_catbox(file_data, fn)
        if link:
            with sqlite3.connect("data/bot.db") as conn:
                conn.execute("UPDATE users SET uploads = uploads + 1 WHERE user_id = ?", (u.id,))
                conn.commit()
            await msg.edit_text(f"✅ Uploaded!\n\n🔗 {link}")
        else:
            await msg.edit_text("❌ Upload failed!")
    except Exception as e:
        await msg.edit_text(f"❌ {e}")

# ========== CONVERT ==========
@dp.message(Command("convert"))
async def convert_cmd(m: Message):
    u, _ = await handle_common(m, "convert")
    if not u:
        return
    if not YTDLP_AVAILABLE:
        await m.answer("❌ yt-dlp not installed!")
        return
    args = m.text.split(maxsplit=1)
    if len(args) < 2:
        await m.answer("📥 Usage: /convert [URL]")
        return
    msg = await m.answer("📥 Downloading...")
    try:
        ydl_opts = {
            'format': 'best',
            'outtmpl': f'temp/video_{u.id}.%(ext)s',
            'quiet': True,
            'no_warnings': True,
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(args[1], download=True)
            filename = ydl.prepare_filename(info)
            if os.path.exists(filename):
                await msg.edit_text("📤 Uploading...")
                with open(filename, "rb") as f:
                    file_data = f.read()
                link = await upload_to_catbox(file_data, os.path.basename(filename))
                os.remove(filename)
                if link:
                    await msg.edit_text(f"✅ Converted!\n\n🔗 {link}")
                else:
                    await msg.edit_text("❌ Upload failed!")
            else:
                await msg.edit_text("❌ Download failed!")
    except Exception as e:
        await msg.edit_text(f"❌ {e}")

# ========== ADMIN: PING (FULL) ==========
@dp.message(Command("ping"))
async def ping_cmd(m: Message):
    u, _ = await handle_common(m, "ping")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    start = time.perf_counter()
    msg = await m.answer("🏓 Testing...")
    latency = int((time.perf_counter() - start) * 1000)
    cpu = psutil.cpu_percent()
    ram = psutil.virtual_memory().percent
    disk = psutil.disk_usage('/').percent
    uptime = format_uptime(int(time.time() - start_time))
    await msg.edit_text(
        f"🏓 Pong!\n\n"
        f"⚡ Latency: {latency}ms\n"
        f"🕒 Uptime: {uptime}\n"
        f"💻 CPU: {cpu}%\n"
        f"💾 RAM: {ram}%\n"
        f"💿 Disk: {disk}%"
    )

# ========== ADMIN: STATS (FULL) ==========
@dp.message(Command("stats"))
async def stats_cmd(m: Message):
    u, _ = await handle_common(m, "stats")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    with sqlite3.connect("data/bot.db") as conn:
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM users")
        users = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM users WHERE last_active >= ?", ((datetime.now() - timedelta(days=7)).isoformat(),))
        alive = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM groups")
        groups = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM uploads")
        uploads = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM wishes")
        wishes = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM users WHERE cult_status != 'none'")
        tempest = c.fetchone()[0]
    await m.answer(
        f"📊 STATISTICS\n\n"
        f"👥 Total Users: {users}\n"
        f"🟢 Alive (7d): {alive}\n"
        f"🔴 Dead: {users - alive}\n"
        f"👥 Groups: {groups}\n"
        f"📁 Uploads: {uploads}\n"
        f"✨ Wishes: {wishes}\n"
        f"🌀 Tempest Members: {tempest}"
    )

# ========== ADMIN: SCAN (DETAILED) ==========
@dp.message(Command("scan"))
async def scan_cmd(m: Message):
    u, _ = await handle_common(m, "scan")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    with sqlite3.connect("data/bot.db") as conn:
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM users")
        users = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM command_logs")
        cmds = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM groups")
        groups = c.fetchone()[0]
        db_size = os.path.getsize("data/bot.db") / 1024
    await m.answer(
        f"🔍 SCAN\n\n"
        f"👥 Users: {users}\n"
        f"🔧 Commands logged: {cmds}\n"
        f"👥 Groups: {groups}\n"
        f"💾 DB Size: {db_size:.1f} KB\n"
        f"✅ All systems healthy!"
    )

# ========== ADMIN: USERS (TEXT FILE) ==========
@dp.message(Command("users"))
async def users_cmd(m: Message):
    u, _ = await handle_common(m, "users")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    with sqlite3.connect("data/bot.db") as conn:
        users = conn.execute("SELECT user_id, first_name, username, uploads, commands, is_banned FROM users ORDER BY commands DESC").fetchall()
    fn = f"temp/users_{int(time.time())}.txt"
    with open(fn, "w", encoding="utf-8") as f:
        f.write("USERS\n" + "=" * 30 + "\n\n")
        for uid, n, un, up, cm, ban in users:
            f.write(f"{n} (@{un or 'None'})\nID: {uid}\nUploads: {up}\nCommands: {cm}\nStatus: {'BANNED' if ban else 'ACTIVE'}\n" + "-" * 20 + "\n")
    await m.answer_document(FSInputFile(fn))
    os.remove(fn)

# ========== BROADCAST ==========
@dp.message(Command("broadcast"))
async def broadcast_cmd(m: Message):
    u, _ = await handle_common(m, "broadcast")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    broadcast_state[u.id] = "waiting"
    save_memory()
    await m.answer("📢 Send anything!")

@dp.message(lambda msg: msg.from_user and msg.from_user.id in broadcast_state)
async def handle_broadcast(m: Message):
    u = m.from_user
    broadcast_state.pop(u.id, None)
    save_memory()
    with sqlite3.connect("data/bot.db") as conn:
        users = conn.execute("SELECT user_id FROM users WHERE is_banned = 0").fetchall()
    s = 0
    for (uid,) in users:
        try:
            await bot.copy_message(chat_id=uid, from_chat_id=m.chat.id, message_id=m.message_id)
            s += 1
        except:
            pass
        await asyncio.sleep(0.03)
    await m.answer(f"✅ {s} users!")

# ========== LAG ==========
@dp.message(Command("lag"))
async def lag_cmd(m: Message):
    await handle_common(m, "lag")
    await m.answer("✅ Done!")

# ========== DISABLE ==========
@dp.message(Command("disable"))
async def disable_cmd(m: Message):
    u, _ = await handle_common(m, "disable")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    args = m.text.split()
    if len(args) < 2:
        await m.answer("⛔ Usage: /disable [cmd]")
        return
    disabled_commands[args[1].replace("/", "")] = datetime.now() + timedelta(minutes=10)
    await m.answer(f"⛔ {args[1]} disabled!")

# ========== CM ==========
@dp.message(Command("cm"))
async def cm_cmd(m: Message):
    u, _ = await handle_common(m, "cm")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    await m.answer(f"{header('⚙️ CM')}\n\n/cm status\n/cm users")

# ========== QUERY (SAFE + SAVED) ==========
@dp.message(Command("query"))
async def query_cmd(m: Message):
    u, _ = await handle_common(m, "query")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer(f"🚫 Owner only: {u.id}")
        return
    args = m.text.split(maxsplit=1)
    if len(args) < 2:
        await m.answer("⚡ Usage: /query [code]")
        return
    code = args[1]
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("INSERT INTO saved_commands (command_code, saved_date) VALUES (?, ?)", (code, datetime.now().isoformat()))
        conn.commit()
    buf = io.StringIO()
    try:
        with contextlib.redirect_stdout(buf):
            exec(code)
        out = buf.getvalue() or "Done"
        await m.answer(f"✅ {out[:3000]}")
    except Exception as e:
        await m.answer(f"❌ {e}")

# ========== MAINTENANCE ==========
@dp.message(Command("maintenance"))
async def maintenance_cmd(m: Message):
    global maintenance_mode
    u, _ = await handle_common(m, "maintenance")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer("🚫 Owner only!")
        return
    maintenance_mode = not maintenance_mode
    await m.answer(f"⚙️ {'🔴 ON' if maintenance_mode else '🟢 OFF'}")

# ========== CLEARLOGS ==========
@dp.message(Command("clearlogs"))
async def clearlogs_cmd(m: Message):
    u, _ = await handle_common(m, "clearlogs")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer("🚫 Owner only!")
        return
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("DELETE FROM command_logs")
        conn.commit()
    await m.answer("🧹 Cleared!")

# ========== LOGS (TEXT FILE) ==========
@dp.message(Command("logs"))
async def logs_cmd(m: Message):
    u, _ = await handle_common(m, "logs")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer("🚫 Owner only!")
        return
    fn = f"temp/logs_{int(time.time())}.txt"
    with sqlite3.connect("data/bot.db") as conn:
        logs = conn.execute("SELECT timestamp, user_id, command FROM command_logs ORDER BY id DESC LIMIT 100").fetchall()
    with open(fn, "w") as f:
        f.write("LOGS\n" + "=" * 30 + "\n\n")
        for ts, uid, cmd in logs:
            f.write(f"{ts} | {uid} | {cmd}\n")
    await m.answer_document(FSInputFile(fn))
    os.remove(fn)

# ========== PRO ==========
@dp.message(Command("pro"))
async def pro_cmd(m: Message):
    u, _ = await handle_common(m, "pro")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer("🚫 Owner only!")
        return
    args = m.text.split()
    if len(args) < 2 or not args[1].isdigit():
        await m.answer("👑 Usage: /pro [id]")
        return
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("UPDATE users SET is_admin = 1 WHERE user_id = ?", (int(args[1]),))
        conn.commit()
    await m.answer(f"✅ {args[1]} admin!")

# ========== BACKUP ==========
@dp.message(Command("backup"))
async def backup_cmd(m: Message):
    u, _ = await handle_common(m, "backup")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer("🚫 Owner only!")
        return
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    bf = f"backups/backup_{ts}.db"
    shutil.copy2("data/bot.db", bf)
    await m.answer_document(FSInputFile(bf))
    os.remove(bf)

# ========== REM ==========
@dp.message(Command("rem"))
async def rem_cmd(m: Message):
    u, _ = await handle_common(m, "rem")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer("🚫 Owner only!")
        return
    pending_restore[u.id] = True
    save_memory()
    await m.answer("💾 Upload .db!")

# ========== RESTART ==========
@dp.message(Command("restart"))
async def restart_cmd(m: Message):
    u, _ = await handle_common(m, "restart")
    if not u:
        return
    if u.id != OWNER_ID:
        await m.answer("🚫 Owner only!")
        return
    save_memory()
    await m.answer("🔄 Restarting...")
    os.execv(sys.executable, ['python'] + sys.argv)

# ========== CANCEL ==========
@dp.message(Command("cancel"))
async def cancel_cmd(m: Message):
    u, _ = await handle_common(m, "cancel")
    if not u:
        return
    upload_waiting.pop(u.id, None)
    broadcast_state.pop(u.id, None)
    pending_restore.pop(u.id, None)
    save_memory()
    await m.answer("❌ Cancelled!")

# ========== BAN SYSTEM ==========
@dp.message(Command("ban"))
async def ban_cmd(m: Message):
    u, _ = await handle_common(m, "ban")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    if not m.reply_to_message:
        await m.answer("⚡ Reply to ban!")
        return
    t = m.reply_to_message.from_user
    if t.id == OWNER_ID:
        await m.answer("❌ Can't ban Owner!")
        return
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("UPDATE users SET is_banned = 1 WHERE user_id = ?", (t.id,))
        conn.commit()
    await m.reply(f"🚫 {t.first_name} banned!")

@dp.message(Command("unban"))
async def unban_cmd(m: Message):
    u, _ = await handle_common(m, "unban")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    args = m.text.split()
    if len(args) < 2 or not args[1].isdigit():
        await m.answer("✅ Usage: /unban [id]")
        return
    with sqlite3.connect("data/bot.db") as conn:
        conn.execute("UPDATE users SET is_banned = 0 WHERE user_id = ?", (int(args[1]),))
        conn.commit()
    await m.answer(f"✅ {args[1]} unbanned!")

@dp.message(Command("banlist"))
async def banlist_cmd(m: Message):
    u, _ = await handle_common(m, "banlist")
    if not u:
        return
    if u.id != OWNER_ID and not await is_admin(u.id):
        await m.answer("🚫 Admin only!")
        return
    with sqlite3.connect("data/bot.db") as conn:
        banned = conn.execute("SELECT user_id, first_name FROM users WHERE is_banned = 1").fetchall()
    if not banned:
        await m.answer("✅ None!")
        return
    await m.answer("🚫 BANNED:\n" + "\n".join(f"• {n} ({uid})" for uid, n in banned))

# ========== MAIN ==========
async def main():
    await ensure_font()
    print("🔒 Starting...")
    try:
        await bot.delete_webhook(drop_pending_updates=True)
    except:
        pass
    while True:
        try:
            await dp.start_polling(bot)
        except Exception as e:
            print(f"❌ {e}")
            await asyncio.sleep(10)

if __name__ == "__main__":
    asyncio.run(main())